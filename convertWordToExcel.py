import docx
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
import tkinter as tk
from tkinter import filedialog
import os

# Set up file dialog
root = tk.Tk()
root.withdraw()  # Hide the main window

# Ask user to select input file
print("Please select the Word document to process...")
INPUT_FILE = filedialog.askopenfilename(
    title="Select Word Document",
    filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
    initialdir="textSources"
)

if not INPUT_FILE:
    print("No file selected. Exiting...")
    exit()

# Create output filename based on input filename
basename = os.path.splitext(os.path.basename(INPUT_FILE))[0]
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)  # Create output directory if it doesn't exist
OUTPUT_FILE = os.path.join(output_dir, f"{basename}.xlsx")

# Function to detect and convert common superscript patterns in text
# Handles:
# 1. French century notation: 'XIXe' -> 'XIX<sup>e</sup>'
# 2. English ordinals: '1st', '2nd', '3rd', '4th' -> '1<sup>st</sup>', etc.
# 3. Numerical measurements: 'km2' -> 'km<sup>2</sup>'
def convert_superscripts(text):
    import re
    # Handle French century numbers (e.g., "XIXe" -> "XIX<sup>e</sup>")
    text = re.sub(r'([XVI]+)e\b', r'\1<sup>e</sup>', text)
    # Handle English ordinals (e.g., "19th" -> "19<sup>th</sup>")
    text = re.sub(r'(\d+)(st|nd|rd|th)\b', r'\1<sup>\2</sup>', text)
    # Handle squared and cubic measurements (e.g., "km2" -> "km<sup>2</sup>")
    text = re.sub(r'([a-zA-Z]+)(\d+)\b', r'\1<sup>\2</sup>', text)
    return text

# Function to process paragraph-level formatting, including:
# - Italics (<em> tags)
# - Line breaks (<br> tags)
# - Explicit superscript from Word document (<sup> tags)
# - Automatic superscript detection for:
#   * French centuries (XIXe -> XIX<sup>e</sup>)
#   * English ordinals (19th -> 19<sup>th</sup>)
#   * Numerical measurements (km2 -> km<sup>2</sup>)
def process_paragraph(paragraph):
    processed_text = ""
    for run in paragraph.runs:
        if run.italic:
            processed_text += f"<em>{run.text}</em>"
        elif (hasattr(run._element.rPr, 'vertAlign') and 
              run._element.rPr.vertAlign is not None and 
              hasattr(run._element.rPr.vertAlign, 'val') and
              run._element.rPr.vertAlign.val == 'superscript'):
            # Handle explicit superscript formatting from Word
            processed_text += f"<sup>{run.text}</sup>"
        else:
            processed_text += run.text
    # Apply automatic superscript detection after handling explicit formatting
    processed_text = convert_superscripts(processed_text)
    return processed_text.replace('\n', '<br>')

# Function to process lists (both numbered and bullet points)
def process_list(paragraph, list_type):
    list_text = ""
    if list_type == "start":
        list_text += "<ul>" if 'Bullet' in paragraph.style.name else "<ol>"
    list_text += f"<li>{process_paragraph(paragraph)}</li>"
    if list_type == "end":
        list_text += "</ul>" if 'Bullet' in paragraph.style.name else "</ol>"
    return list_text



def micro_typographic_rules(text):
    """
    Apply micro-typographic rules: replace spaces in various typographic situations with HTML non-breaking spaces (&nbsp;).
    - After '«' and before '»'
    - Before semicolon, colon, and en dash (–)
    - After ordinal superscripts only (XIX<sup>e</sup>, 19<sup>th</sup>)
    - Around measurement units (km, m, cm, etc.)
    - Around multiplication symbols (× or x between numbers)
    - Before currency symbols (€, $, £)
    - Before percentage and degree symbols (%, °)
    Example: '« text » ; : – XIX<sup>e</sup> siècle 20 km × 30 m 25 °C 100 €' 
        -> '«&nbsp;text&nbsp;»&nbsp;;&nbsp;:&nbsp;–&nbsp;XIX<sup>e</sup>&nbsp;siècle 20&nbsp;km&nbsp;×&nbsp;30&nbsp;m 25&nbsp;°C 100&nbsp;€'
    """
    import re
    # Replace space after « with &nbsp;
    text = re.sub(r'«\s+', '«&nbsp;', text)
    # Replace space before » with &nbsp;
    text = re.sub(r'\s+»', '&nbsp;»', text)
    # Replace space before semicolon with &nbsp;
    text = re.sub(r'\s+;', '&nbsp;;', text)
    # Replace space before colon with &nbsp;
    text = re.sub(r'\s+:', '&nbsp;:', text)
    # Replace space before en dash (–) with &nbsp;
    text = re.sub(r'\s+–', '&nbsp;–', text)
    # Replace space only after ordinal superscripts (French e/er/ère and English th/st/nd/rd)
    text = re.sub(r'<sup>(e|er|ère|th|st|nd|rd)</sup>\s+', r'<sup>\1</sup>&nbsp;', text)
    # Measurement units and dimensions (English, French, German)
    units = r'(?:km|m|cm|mm|ha|mètres?|meters?|Meter|Metern)'
    text = re.sub(rf'(\d+)\s+({units})\b', r'\1&nbsp;\2', text)
    # Multiplication symbol in dimensions (only between numbers with optional spaces)
    text = re.sub(r'(\d+)\s*([×x])\s*(\d+)', r'\1&nbsp;\2&nbsp;\3', text)
    # Temperature, percentages, currencies, and degrees (ensuring symbols don't appear mid-word)
    text = re.sub(r'(\d+)\s*([°%€$£]|EUR|CHF|USD)\b', r'\1&nbsp;\2', text)
    # Time units (avoiding line breaks between number and unit)
    time_units = r'(?:h|min|s|Uhr|heures?|hours?|Stunden?)'
    text = re.sub(rf'(\d+)\s+({time_units})\b', r'\1&nbsp;\2', text)
    # Common abbreviations that shouldn't break (ensuring we match complete abbreviations)
    text = re.sub(r'\b(z\.\s*B\.|d\.\s*h\.|i\.\s*e\.|e\.\s*g\.|etc\.)\s+', r'\1&nbsp;', text)
    return text

# Load the Word document
doc = docx.Document(INPUT_FILE)

# Initialize an empty list to store the rows for the Excel file
rows = []

# Iterate through each table in the document
for table in doc.tables:
    # Initialize a dictionary to store the row data
    row_data = {}
    # Iterate through each row in the table
    for row in table.rows:
        # The first cell is the column name, the second cell is the value
        column_name = row.cells[0].text.strip()
        value = ""
        list_open = False
        for i, paragraph in enumerate(row.cells[1].paragraphs):
            if paragraph.style.name.startswith('List'):
                if not list_open:
                    value += process_list(paragraph, "start")
                    list_open = True
                value += f"<li>{process_paragraph(paragraph)}</li>"
            else:
                if list_open:
                    value += process_list(paragraph, "end")
                    list_open = False
                value += process_paragraph(paragraph)
                if i < len(row.cells[1].paragraphs) - 1:
                    value += "<br>"
        if list_open:
            value += process_list(paragraph, "end")
            list_open = False  # Ensure list is closed
        # Add the data to the dictionary
        if column_name in row_data:
            row_data[column_name] += "<br>" + value
        else:
            row_data[column_name] = value
    # Append the dictionary to the list of rows
    rows.append(row_data)


# Apply micro-typographic rules to all fields in all rows
for row in rows:
    for key in row:
        row[key] = micro_typographic_rules(row[key])

# Convert the list of dictionaries to a DataFrame
df = pd.DataFrame(rows)

# Save the DataFrame to an Excel file
df.to_excel(OUTPUT_FILE, index=False)

# Load the workbook and select the active worksheet
wb = load_workbook(OUTPUT_FILE)
ws = wb.active

# Set the width of the column with title 'Fliesstext' to 400px and wrap its contents
for col in ws.iter_cols(1, ws.max_column):
    if col[0].value == 'Fliesstext':
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = 400 / 7  # Convert pixels to Excel's column width units
        for cell in col:
            cell.alignment = Alignment(wrap_text=True)

# Save the updated workbook
wb.save(OUTPUT_FILE)

print(f"The tables have been successfully combined into '{OUTPUT_FILE}' with adjusted column width and wrapping.")